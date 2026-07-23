import streamlit as st
import google.generativeai as genai
import datetime
import io
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from supabase import create_client, Client
from PIL import Image

st.set_page_config(page_title="Nyaya Assist AI", page_icon="⚖️", layout="wide", initial_sidebar_state="expanded")

TEMPLATES_FOLDER = "master_templates"
if not os.path.exists(TEMPLATES_FOLDER):
    os.makedirs(TEMPLATES_FOLDER)

# 🔑 Fetch Secrets
default_api_key = st.secrets.get("GEMINI_API_KEY", "")
supabase_url = st.secrets.get("SUPABASE_URL", "")
supabase_key = st.secrets.get("SUPABASE_KEY", "")

# 🗄️ Supabase Initialization
supabase: Client = None
if supabase_url and supabase_key:
    try:
        supabase = create_client(supabase_url, supabase_key)
    except Exception as e:
        st.error(f"Supabase connection error: {e}")

# 🔐 PREMIUM MODERN DARK CSS STYLING
st.markdown("""
<style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    .stDeployButton {display:none;}
    
    .stApp {
        background-color: #050508 !important;
        color: #f3f4f6 !important;
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
    }
    
    section[data-testid="stSidebar"] {
        background-color: #08090d !important;
        border-right: 1px solid #161821 !important;
        padding-top: 10px;
    }
    
    .stTextInput > div > div > input, .stSelectbox > div > div, .stTextArea > div > div > textarea {
        background-color: #0e1017 !important;
        color: #ffffff !important;
        border: 1px solid #1f222e !important;
        border-radius: 12px !important;
    }
    
    .stButton > button {
        background: linear-gradient(135deg, #7c3aed, #6366f1) !important;
        color: white !important;
        border-radius: 10px !important;
        border: none !important;
        padding: 8px 16px !important;
        font-weight: 600 !important;
        box-shadow: 0 4px 12px rgba(124, 58, 237, 0.3) !important;
        width: 100%;
    }
    
    .stDownloadButton > button {
        background: #1e1b4b !important;
        color: #c084fc !important;
        border: 1px solid #a855f7 !important;
        border-radius: 10px !important;
        width: 100%;
    }

    .pro-banner {
        background: linear-gradient(135deg, #1e1b4b, #0f172a);
        border: 1px solid #3b0764;
        border-radius: 14px;
        padding: 12px;
        margin-bottom: 20px;
    }

    .welcome-hero {
        text-align: center;
        padding: 10px;
    }
    .welcome-title {
        font-size: 28px;
        font-weight: 700;
        color: #ffffff;
    }
    .welcome-subtitle {
        color: #9ca3af;
        font-size: 14px;
        margin-bottom: 20px;
    }
    
    .case-studio-banner {
        background-color: #0e1017;
        border: 1px solid #1f222e;
        border-radius: 16px;
        padding: 16px;
        margin: 0 auto 15px auto;
    }

    .action-card {
        background-color: #0e1017;
        border: 1px solid #1f222e;
        border-radius: 12px;
        padding: 12px;
        margin-bottom: 8px;
    }

    .attachment-chip {
        background-color: #1e1b4b;
        border: 1px solid #7c3aed;
        border-radius: 10px;
        padding: 8px 12px;
        color: #c084fc;
        font-size: 13px;
        font-weight: 600;
        margin-bottom: 10px;
    }
</style>
""", unsafe_allow_html=True)

if "user" not in st.session_state:
    st.session_state.user = None
if "auth_mode" not in st.session_state:
    st.session_state.auth_mode = "login"
if "uploader_key" not in st.session_state:
    st.session_state.uploader_key = 0
if "current_session" not in st.session_state:
    st.session_state.current_session = f"Case_{datetime.datetime.now().strftime('%d%b_%H%M')}"
if "nav_menu" not in st.session_state:
    st.session_state.nav_menu = "💬 Case Studio"
if "show_profile" not in st.session_state:
    st.session_state.show_profile = False

params = st.query_params
if "logged_in_user" in params and not st.session_state.user:
    saved_email = params["logged_in_user"]
    class PersistentUser:
        def __init__(self, email):
            self.email = email
            self.user_metadata = {"full_name": email.split('@')[0].capitalize()}
    st.session_state.user = PersistentUser(saved_email)

def login_user(email, password):
    try:
        res = supabase.auth.sign_in_with_password({"email": email, "password": password})
        st.session_state.user = res.user
        st.query_params["logged_in_user"] = email
        st.success("✅ Login successful!")
        st.rerun()
    except Exception as e:
        st.error(f"Login failed: {e}")

def signup_user(email, password, full_name, phone):
    try:
        res = supabase.auth.sign_up({
            "email": email, "password": password,
            "options": {"data": {"full_name": full_name, "phone": phone}}
        })
        st.session_state.user = res.user
        st.query_params["logged_in_user"] = email
        st.success("✅ Account created successfully!")
        st.rerun()
    except Exception as e:
        st.error(f"Signup failed: {e}")

if not st.session_state.user:
    st.markdown("<br>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1, 4, 1])
    with col2:
        if st.session_state.auth_mode == "login":
            st.markdown("<h2 style='text-align:center; color:white;'>⚖️ Nyaya Assist AI</h2>", unsafe_allow_html=True)
            email = st.text_input("Email", key="l_email")
            password = st.text_input("Password", type="password", key="l_pass")
            if st.button("Login", use_container_width=True):
                login_user(email, password)
            if st.button("Create account", use_container_width=True):
                st.session_state.auth_mode = "signup"
                st.rerun()
        else:
            st.markdown("<h2 style='color:white;'>Create account</h2>", unsafe_allow_html=True)
            fname = st.text_input("Full Name", key="s_name")
            email = st.text_input("Email", key="s_email")
            phone = st.text_input("Phone", key="s_phone")
            password = st.text_input("Password", type="password", key="s_pass")
            if st.button("Sign Up", use_container_width=True):
                signup_user(email, password, fname, phone)
            if st.button("Already have an account? Login", use_container_width=True):
                st.session_state.auth_mode = "login"
                st.rerun()
    st.stop()

current_user_email = st.session_state.user.email
display_user_name = current_user_email.split('@')[0].capitalize()
user_initials = "".join([part[0].upper() for part in display_user_name.split()[:2]])

def get_supabase_sessions():
    if supabase:
        try:
            res = supabase.table("chats").select("session_name").eq("user_email", current_user_email).execute()
            return sorted(list(set([row["session_name"] for row in res.data])), reverse=True)
        except Exception:
            return []
    return []

def get_supabase_chat_history(session_name):
    if supabase:
        try:
            res = supabase.table("chats").select("role, content").eq("user_email", current_user_email).eq("session_name", session_name).order("id").execute()
            return res.data
        except Exception:
            return []
    return []

def save_chat_to_supabase(session_name, role, content):
    if supabase:
        try:
            supabase.table("chats").insert({"user_email": current_user_email, "session_name": session_name, "role": role, "content": content}).execute()
        except Exception:
            pass

def delete_supabase_session(session_name):
    if supabase:
        try:
            supabase.table("chats").delete().eq("user_email", current_user_email).eq("session_name", session_name).execute()
            st.rerun()
        except Exception:
            pass

def read_uploaded_file_content(f):
    if f is None: return "", None
    ft = f.name.split('.')[-1].lower()
    if ft == "txt": return f.read().decode("utf-8"), None
    elif ft == "docx": return '\n'.join([p.text for p in Document(f).paragraphs]), None
    elif ft == "pdf":
        try:
            import pypdf
            return "\n".join([pg.extract_text() for pg in pypdf.PdfReader(f).pages if pg.extract_text()]), None
        except: return "[PDF]", None
    elif ft in ["jpg", "jpeg", "png"]: return f"[IMAGE]", Image.open(f)
    return "", None

def create_court_ready_docx(text):
    doc = Document()
    p = doc.add_paragraph(text)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

nav_options = ["💬 Case Studio", "📖 Library", "📁 My Cases", "⚙️ Settings", "📂 Chat History"]

# 🎯 SIDEBAR NAVIGATION
with st.sidebar:
    st.markdown("""
    <div style="text-align: center; margin-bottom: 15px;">
        <h3>⚖️ Nyaya Assist AI</h3>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("""
    <div class="pro-banner">
        <div style="font-weight: 700; font-size: 12px; color: #ffffff;">👑 Upgrade to Pro</div>
        <div style="font-size: 10px; color: #9ca3af;">Unlock advanced features</div>
    </div>
    """, unsafe_allow_html=True)

    if st.button("➕ New Case", key="sb_new", use_container_width=True):
        st.session_state.current_session = f"Case_{datetime.datetime.now().strftime('%d%b_%H%M')}"
        st.session_state.nav_menu = "💬 Case Studio"
        st.rerun()

    curr_idx = nav_options.index(st.session_state.nav_menu) if st.session_state.nav_menu in nav_options else 0
    menu = st.radio("Menu", nav_options, index=curr_idx, label_visibility="collapsed")
    st.session_state.nav_menu = menu

    st.markdown("---")
    st.markdown("<div style='font-size: 11px; font-weight:700; color:#8e92a4; text-transform:uppercase;'>Chat History</div>", unsafe_allow_html=True)
    for s in get_supabase_sessions()[:10]:
        if st.button(f"💬 {s[:20]}...", key=f"sbar_{s}", use_container_width=True):
            st.session_state.current_session = s
            st.session_state.nav_menu = "💬 Case Studio"
            st.rerun()

    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("🚪 Logout", use_container_width=True):
        if "logged_in_user" in st.query_params: del st.query_params["logged_in_user"]
        st.session_state.user = None
        st.rerun()

# TOP APP BAR WITH WORKING PROFILE TOGGLE
col1, col2, col3 = st.columns([5, 1, 1])
with col1: st.markdown("### Nyaya Assist <span style='color:#a855f7;'>AI</span>", unsafe_allow_html=True)
with col3: 
    if st.button(user_initials, key="prof_btn"):
        st.session_state.show_profile = not st.session_state.show_profile

# 👤 PROFILE PANEL TOGGLE VIEW
if st.session_state.show_profile:
    st.markdown(f"""
    <div style="background-color: #0e1017; border: 1px solid #7c3aed; border-radius: 12px; padding: 15px; margin-bottom: 20px;">
        <h4 style="margin: 0; color: #ffffff;">👤 Account Profile</h4>
        <p style="color: #9ca3af; margin: 5px 0;"><b>Email:</b> {current_user_email}</p>
        <p style="color: #9ca3af; margin: 5px 0;"><b>Plan:</b> Free Tier 👑</p>
    </div>
    """, unsafe_allow_html=True)

if default_api_key:
    genai.configure(api_key=default_api_key)
    model = genai.GenerativeModel(model_name="gemini-3.6-flash", generation_config={"temperature": 0.0})

    if menu == "📖 Library":
        st.markdown("## 📖 Legal Library")
        q = st.text_input("Search Library...")
        up = st.file_uploader("Upload File", type=["pdf","docx","txt"])
        if up:
            with open(os.path.join(TEMPLATES_FOLDER, up.name), "wb") as f: f.write(up.getbuffer())
            st.success("Uploaded!")
        for f in os.listdir(TEMPLATES_FOLDER):
            st.markdown(f"📄 `{f}`")

    elif menu == "📁 My Cases":
        st.markdown("## 📁 My Cases")
        for s in get_supabase_sessions(): st.markdown(f"- 📁 `{s}`")

    elif menu == "⚙️ Settings":
        st.markdown("## ⚙️ Settings")
        adv = st.text_input("Advocate Name", value=st.session_state.get("advocate_name",""))
        court = st.text_input("Court", value=st.session_state.get("default_court",""))
        if st.button("Save"):
            st.session_state.advocate_name = adv
            st.session_state.default_court = court
            st.success("Saved!")

    elif menu == "📂 Chat History":
        st.markdown("## 📂 Chat Archives")
        for s in get_supabase_sessions():
            if st.button(f"Open {s}", key=f"ch_{s}"):
                st.session_state.current_session = s
                st.session_state.nav_menu = "💬 Case Studio"
                st.rerun()

    elif menu == "💬 Case Studio":
        messages = get_supabase_chat_history(st.session_state.current_session)
        if not messages:
            st.markdown(f"""
            <div class="welcome-hero">
                <div class="welcome-title">Welcome back, <span style="color: #a855f7;">{display_user_name}</span></div>
                <div class="welcome-subtitle">Your AI Legal Assistant is ready to support your legal work.</div>
            </div>
            """, unsafe_allow_html=True)
            if st.button("➕ Start New Case", use_container_width=True):
                st.session_state.current_session = f"Case_{datetime.datetime.now().strftime('%d%b_%H%M')}"
                st.rerun()
        else:
            for idx, msg in enumerate(messages):
                with st.chat_message(msg["role"]):
                    st.markdown(msg["content"])
                    if msg["role"] == "assistant" and "START_DRAFT" in msg["content"].upper():
                        st.download_button("📥 Download Word (.docx)", create_court_ready_docx(msg["content"]), file_name=f"Draft_{idx}.docx", key=f"dw_{idx}")

        # INPUT DECK WITH TEMPLATE, CASE DOCS & VOICE
        st.markdown("---")
        c1, c2, c3 = st.columns(3)
        with c1:
            st.markdown("<div class='action-card'><strong style='color:#c084fc; font-size:11px;'>📤 Template</strong></div>", unsafe_allow_html=True)
            f_file = st.file_uploader("Template", type=["docx","txt","pdf"], key=f"f_{st.session_state.uploader_key}", label_visibility="collapsed")
        with c2:
            st.markdown("<div class='action-card'><strong style='color:#c084fc; font-size:11px;'>📸 Case Docs</strong></div>", unsafe_allow_html=True)
            c_file = st.file_uploader("Case", type=["docx","txt","pdf","jpg","png"], key=f"c_{st.session_state.uploader_key}", label_visibility="collapsed")
        with c3:
            st.markdown("<div class='action-card'><strong style='color:#c084fc; font-size:11px;'>🎙️ Voice Assistant</strong></div>", unsafe_allow_html=True)
            voice_html = """
            <div style="text-align: center;">
                <button id="recordBtn" onclick="toggleRecord()" style="background: linear-gradient(135deg, #7c3aed, #6366f1); color: white; border: none; padding: 6px 8px; border-radius: 6px; font-weight: 600; cursor: pointer; font-size: 11px; width: 100%;">🎤 Speak</button>
                <div id="statusTxt" style="color: #a855f7; font-size: 9px; margin-top: 2px;"></div>
                <textarea id="speechOutput" placeholder="Voice text..." style="width: 100%; height: 26px; background-color: #050508; color: #ffffff; border: 1px solid #1f222e; border-radius: 4px; padding: 2px; font-size: 10px; resize: none; margin-top: 2px;"></textarea>
            </div>
            <script>
                var recognition; var isRecording = false;
                function toggleRecord() {
                    if (!('webkitSpeechRecognition' in window) && !('SpeechRecognition' in window)) { alert("Use Chrome."); return; }
                    if (!isRecording) {
                        var SpeechRecognition = window.SpeechRecognition || window.webkitSpeechRecognition;
                        recognition = new SpeechRecognition(); recognition.continuous = true; recognition.interimResults = true; recognition.lang = 'hi-IN';
                        recognition.onstart = function() { isRecording = true; document.getElementById('recordBtn').innerText = "🛑 Stop"; document.getElementById('recordBtn').style.background = "#ef4444"; };
                        recognition.onresult = function(event) {
                            var t = ''; for (var i = event.resultIndex; i < event.results.length; ++i) { t += event.results[i][0].transcript; }
                            document.getElementById('speechOutput').value = t;
                        };
                        recognition.onend = function() { isRecording = false; document.getElementById('recordBtn').innerText = "🎤 Speak"; document.getElementById('recordBtn').style.background = "linear-gradient(135deg, #7c3aed, #6366f1)"; };
                        recognition.start();
                    } else { recognition.stop(); }
                }
            </script>
            """
            st.components.v1.html(voice_html, height=75)

        user_input = st.chat_input("Ask Nyaya AI...")
        if user_input or f_file or c_file:
            fc, _ = read_uploaded_file_content(f_file) if f_file else ("", None)
            cc, c_img = read_uploaded_file_content(c_file) if c_file else ("", None)
            prompt = f"Template: {fc}\nEvidence: {cc}\nInstruction: {user_input or 'Analyze case'}"
            save_chat_to_supabase(st.session_state.current_session, "user", user_input or "Uploaded files")
            with st.chat_message("user"): st.markdown(user_input or "Uploaded files")
            with st.chat_message("assistant"):
                resp = model.generate_content([prompt, c_img] if c_img else [prompt]).text
                st.markdown(resp)
                save_chat_to_supabase(st.session_state.current_session, "assistant", resp)
            st.session_state.uploader_key += 1
            st.rerun()
else:
    st.info("👈 API Key missing.")
